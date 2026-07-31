try:
    from ._imports import without_local_module_shadowing
except ImportError:
    from _imports import without_local_module_shadowing

with without_local_module_shadowing(__file__):
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

try:
    from ._layout import render_grid
    from .source import as_source
except ImportError:
    from _layout import render_grid
    from source import as_source


class DOCXExtractor(object):
    def __init__(self, filename, include_headers=True, name=None):
        self.source = as_source(filename, name)
        self.filename = self.source.name
        self.include_headers = include_headers
        self._list_counters = {}
        self.blocks = []
        self.tables = []
        self._extract()
        self.text = self._render()
        self.provenance = self.source.provenance("docx")

    def get_text(self):
        return self.text

    def get_structure(self):
        return {
            "type": "docx",
            "provenance": self.provenance,
            "blocks": self.blocks,
            "tables": self.tables,
        }

    def get_tables(self):
        return self.tables

    def _extract(self):
        with self.source.open() as stream:
            document = Document(stream)
            for item in document.iter_inner_content():
                if isinstance(item, Paragraph):
                    block = self._paragraph_block(item)
                    if block["text"] or block["kind"] != "paragraph":
                        self.blocks.append(block)
                elif isinstance(item, Table):
                    table = self._table_data(item)
                    self.tables.append(table)
                    self.blocks.append({"kind": "table", "table": table})

            if self.include_headers:
                self._extract_headers_and_footers(document)

    def _paragraph_block(self, paragraph):
        text = paragraph.text
        style = paragraph.style.name if paragraph.style else None
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        level = 0
        marker = ""
        style_lower = (style or "").lower()
        is_list_style = "list bullet" in style_lower or "list number" in style_lower
        if num_pr is not None or is_list_style:
            if num_pr is None:
                level = max(style_lower.rsplit(" ", 1)[-1].isdigit() and int(
                    style_lower.rsplit(" ", 1)[-1]
                ) - 1 or 0, 0)
                num_id = style_lower
            else:
                level = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
                num_id = int(num_pr.numId.val) if num_pr.numId is not None else 0
            key = (num_id, level)
            self._list_counters[key] = self._list_counters.get(key, 0) + 1
            for other_key in list(self._list_counters):
                if other_key[0] == num_id and other_key[1] > level:
                    del self._list_counters[other_key]
            marker = (
                "• "
                if "bullet" in style_lower
                else "{}. ".format(self._list_counters[key])
            )
        kind = "heading" if (style or "").lower().startswith("heading") else "paragraph"
        return {
            "kind": kind,
            "text": text,
            "style": style,
            "level": level,
            "marker": marker,
        }

    def _table_data(self, table):
        row_count = len(table.rows)
        col_count = len(table.columns)
        cells = []
        seen = set()
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                identity = cell._tc
                if identity in seen:
                    continue
                seen.add(identity)
                col_span = max(int(cell._tc.tcPr.gridSpan.val), 1) if (
                    cell._tc.tcPr is not None
                    and cell._tc.tcPr.gridSpan is not None
                ) else 1
                row_span = 1
                for next_row in range(row_index + 1, row_count):
                    if col_index >= len(table.rows[next_row].cells):
                        break
                    if table.rows[next_row].cells[col_index]._tc is cell._tc:
                        row_span += 1
                    else:
                        break
                paragraphs = []
                nested_tables = []
                for item in cell.iter_inner_content():
                    if isinstance(item, Paragraph):
                        if item.text:
                            paragraphs.append(item.text)
                    elif isinstance(item, Table):
                        nested = self._table_data(item)
                        nested_tables.append(nested)
                        paragraphs.append(self._render_table(nested))
                cells.append(
                    {
                        "row": row_index,
                        "col": col_index,
                        "rowSpan": row_span,
                        "colSpan": col_span,
                        "text": "\n".join(paragraphs),
                        "nested": nested_tables,
                    }
                )
        widths = [
            float(column.width or 1)
            for column in table.columns
        ]
        return {
            "rows": row_count,
            "cols": col_count,
            "cells": cells,
            "columnWidths": widths,
        }

    def _extract_headers_and_footers(self, document):
        seen_parts = set()
        for section_index, section in enumerate(document.sections):
            for label, container in (
                ("header", section.header),
                ("firstPageHeader", section.first_page_header),
                ("evenPageHeader", section.even_page_header),
                ("footer", section.footer),
                ("firstPageFooter", section.first_page_footer),
                ("evenPageFooter", section.even_page_footer),
            ):
                part_name = str(container.part.partname)
                if part_name in seen_parts:
                    continue
                seen_parts.add(part_name)
                content = []
                for item in container.iter_inner_content():
                    if isinstance(item, Paragraph) and item.text:
                        content.append(item.text)
                    elif isinstance(item, Table):
                        table = self._table_data(item)
                        self.tables.append(table)
                        content.append(self._render_table(table))
                if content:
                    self.blocks.append(
                        {
                            "kind": label,
                            "section": section_index,
                            "text": "\n".join(content),
                        }
                    )

    @staticmethod
    def _render_table(table):
        grid = [["" for _ in range(table["cols"])] for _ in range(table["rows"])]
        for cell in table["cells"]:
            if cell["row"] >= table["rows"] or cell["col"] >= table["cols"]:
                continue
            value = cell["text"]
            if cell["rowSpan"] > 1 or cell["colSpan"] > 1:
                value += " (병합 {}×{})".format(cell["rowSpan"], cell["colSpan"])
            grid[cell["row"]][cell["col"]] = value
        return render_grid(
            grid,
            title="[표 {}×{}]".format(table["rows"], table["cols"]),
            column_weights=table["columnWidths"],
        )

    def _render(self):
        output = []
        labels = {
            "header": "머리말",
            "firstPageHeader": "첫 페이지 머리말",
            "evenPageHeader": "짝수 페이지 머리말",
            "footer": "꼬리말",
            "firstPageFooter": "첫 페이지 꼬리말",
            "evenPageFooter": "짝수 페이지 꼬리말",
        }
        for block in self.blocks:
            if block["kind"] == "table":
                output.append(self._render_table(block["table"]))
            elif block["kind"] in labels:
                output.append("[{}]\n{}".format(labels[block["kind"]], block["text"]))
            else:
                indent = "  " * block.get("level", 0)
                output.append(indent + block.get("marker", "") + block["text"])
        return "\n\n".join(output)


def get_text(filename):
    print(DOCXExtractor(filename).get_text())


if __name__ == "__main__":
    import sys

    get_text(sys.argv[1])
